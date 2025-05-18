import os
import uuid
import PyPDF2
from google.cloud import vision
import docx
import pandas as pd
import io
from typing import Optional, List
from services.read_pdf import AzureAIPDFReader
from config import settings
from models import ItemDetail, FileType
from services.content_verifier import ContentVerifier

class DocumentProcessor:
    def __init__(self):
        """Initialize document processor with necessary clients and configurations"""
        self.vision_client = vision.ImageAnnotatorClient()
        self.azure_endpoint = settings.AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT
        self.azure_key = settings.AZURE_DOCUMENT_INTELLIGENCE_KEY

    def extract_content(self, file_path: str) -> Optional[str]:
        """
        Main method to extract text content from various file formats
        
        Args:
            file_path: Path to the input file
            
        Returns:
            Extracted text content as string, or None if extraction fails
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = self._determine_file_type(file_path)
        
        try:
            if file_type == 'pdf':
                return self._extract_pdf(file_path)
            elif file_type == 'docx':
                return self._extract_docx(file_path)
            elif file_type == 'excel':
                return self._extract_excel(file_path)
            elif file_type == 'image':
                return self._extract_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return None

    def _determine_file_type(self, file_path: str) -> str:
        """Determine file type based on extension"""
        ext = os.path.splitext(file_path)[1].lower().lstrip('.')
        if ext == 'pdf':
            return 'pdf'
        elif ext == 'docx':
            return 'docx'
        elif ext in ['xls', 'xlsx', 'csv']:
            return 'excel'
        elif ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
            return 'image'
        else:
            return ext

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using Azure AI or PyPDF2 fallback"""
        if self.azure_endpoint and self.azure_key:
            try:
                return self._extract_pdf_azure(file_path)
            except Exception as e:
                print(f"Azure extraction failed: {e}, using fallback")
        return self._extract_pdf_fallback(file_path)

    def _extract_pdf_azure(self, file_path: str) -> str:
        """Extract text using Azure Document Intelligence"""
        pdf_reader = AzureAIPDFReader(
            endpoint=self.azure_endpoint,
            key=self.azure_key
        )
        result = pdf_reader.extract_text(file_path)
        text = ""
        
        if result.get("paragraphs"):
            text = "\n\n".join(p["text"] for p in result["paragraphs"])
        elif result.get("pages"):
            for page in result["pages"]:
                text += "\n".join(page.get("lines", [])) + "\n\n"
        
        if result.get("tables"):
            text += "\n\nTABLES:\n"
            for table in result["tables"]:
                if table.get("grid"):
                    for row in table["grid"]:
                        text += " | ".join(str(cell) for cell in row) + "\n"
                    text += "\n"
        return text

    def _extract_pdf_fallback(self, file_path: str) -> str:
        """Fallback PDF extraction using PyPDF2"""
        text = ""
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n\n"
        return text

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX documents"""
        text = []
        doc = docx.Document(file_path)
        
        for para in doc.paragraphs:
            text.append(para.text)
            
        for table in doc.tables:
            for row in table.rows:
                text.append(" | ".join(cell.text for cell in row.cells))
                
        return "\n".join(text)

    def _extract_excel(self, file_path: str) -> str:
        """Extract text from Excel files"""
        df = pd.read_excel(file_path)
        return df.to_csv(index=False)

    def _extract_image(self, file_path: str) -> str:
        """Extract text from images using Google Vision OCR"""
        with open(file_path, "rb") as image_file:
            content = image_file.read()

        image = vision.Image(content=content)
        response = self.vision_client.text_detection(image=image)
        
        if response.error.message:
            raise Exception(f"OCR Error: {response.error.message}")
            
        return response.text_annotations[0].description if response.text_annotations else ""

async def process_document(file_path: str, file_type: FileType) -> List[ItemDetail]:
    """
    Process a document to extract items
    
    Args:
        file_path: Path to the file
        file_type: Type of the file
        
    Returns:
        List of extracted item details
    """
    try:
        # Initialize the document processor
        processor = DocumentProcessor()
        
        # Extract text content
        extracted_text = processor.extract_content(file_path)
        
        if not extracted_text:
            print(f"No text could be extracted from {file_path}")
            return []
        
        # Use content verifier to extract items from the text
        content_verifier = ContentVerifier()
        items, is_rfq = content_verifier.generate_rfq(extracted_text)
        
        return items
    except Exception as e:
        print(f"Error in process_document: {str(e)}")
        raise Exception(f"Error processing document: {str(e)}")